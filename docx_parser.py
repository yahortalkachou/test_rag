import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support disabled.")

class ProjectBlockParser:
    """Parser for detecting and extracting projects block from CV text."""
    
    # Patterns for project block detection
    PROJECT_SECTION_PATTERNS = [
        r'projects?[\s]*:?',
        r'experience[\s]*:?',
        r'work experience[\s]*:?',
        r'project experience[\s]*:?',
        r'проекты?[\s]*:?',
        r'опыт работы[\s]*:?',
    ]
    
    # Patterns for project roles and period
    PROJECT_ROLE_PATTERNS = [
        r'project roles?',
        r'role[s]?',
        r'должность',
        r'роль'
    ]
    
    PROJECT_PERIOD_PATTERNS = [
        r'period',
        r'duration',
        r'время',
        r'период'
    ]
    
    INNOWISE_PROJECTS_SECTION_PATTERNS = [
        
    ]
    
    def __init__(self):
        self.compiled_section_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in self.PROJECT_SECTION_PATTERNS]
        self.compiled_role_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in self.PROJECT_ROLE_PATTERNS]
        self.compiled_period_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in self.PROJECT_PERIOD_PATTERNS]
    
    def _file_check(self, file_path: str):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    def _get_instance_between_keywords (self, lst: List[str], begin: str, end: str) -> List[str]:
        try:
            start = lst.index(begin) + 1
        except Exception as e:
            return []
        try:
            stop = lst.index(end)
            return lst[start:stop]
        except Exception as e:
            return lst[start:]
    
    def _get_innowise_project_from_table (self, cells: List)->Dict:
        descriprion = cells[1].text
        project_name = cells[0].text.split('\n')
        summury = project_name[1]
        project_name = project_name[0]
        specialisations = self._get_instance_between_keywords(descriprion.split('\n'),"Project roles","Period")[0].split('/')
        descriprion = summury + descriprion
        
        return {
            "name": project_name,
            "description": descriprion,
            "specialisations" : specialisations
        }
            
        
        
        
    def parse_inniwise_std (self, file_path) -> Dict[str,any]:
        """
        Extract personal and project data from .docx innowise standard file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Extracted text as string
        """
        self._file_check(file_path)
        name = ""
        specialisations = []
        eduction = ""
        languages = []
        domains = []
        personal_data = ""
        projects = []
        level = ""
        
        para_texts = []
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_texts.append(paragraph.text)
            if len(para_texts) == 3:
                name = para_texts[0]
                level = para_texts[1].split(" ")[0]
                specialisations = para_texts[1].split("/")
                specialisations[0] = " ".join(specialisations[0].split(" ")[1:])
            print(name,level, specialisations)
                            
            # Extract from std tables
            if len(doc.tables) == 3:
                #pesonal/global metadat
                about = doc.tables[0].rows[0].cells[0].text
                about = about.split('\n')
                print (about)
                eduction = self._get_instance_between_keywords(about,"Education","Language proficiency")
                languages = self._get_instance_between_keywords(about,"Language proficiency","Domains")
                domains = self._get_instance_between_keywords(about,"Domains","Certificates")
                cert = self._get_instance_between_keywords(about,"Certificates","axaxaxax")
                print (f"ed: {eduction} dom: {domains} langs: {languages} certs: {cert}")
                description = doc.tables[0].rows[0].cells[1].text.split('\n')[1]
                print(f"description\n{description}")
                #projectdata section
                
                if len(doc.tables) == 3:
                    table = doc.tables[1]
                    for row in table.rows[1:]:
                        projects.append(self._get_innowise_project_from_table(row.cells))
                
            global_data = {
                "name": name,
                "level": level,
                "specialisations": specialisations,
                "education": eduction,
                "languages": languages,
                "domains": domains,
                "projects": projects
            }
            for p in projects:
                print (f"name: {p['name']}\nspecs: {p['specialisations']}\ndescr: {p['description']}\n")
            return global_data
        
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {e}")
        
    def parse_docx(self, file_path: str) -> str:
        """
        Extract text from .docx file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Extracted text as string
        """

        self._file_check(file_path)
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                
                
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    print(f"para\n{paragraph.text}\n")
            
            # Extract text from tables
            for table in doc.tables:
                print(f"\n!!NEW_TABLE!!\n")
                for row in table.rows:
                    print("\n\t!New_Row\n")
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                            print(f"|_\t\t{cell.text}_|\n")
                            
                    if row_text:
                        text_content.append(' | '.join(row_text))
                    #print(f"table\n____________________{''.join(row_text)}\n___________________________\n")
            return '\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {e}")
    
    def find_projects_block_from_docx(self, file_path: str) -> Dict[str, any]:
        """
        Find projects block directly from .docx file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Dict with parsing results
        """
        cv_text = self.parse_docx(file_path)
        return self.find_projects_block(cv_text)
    
    def find_projects_block(self, cv_text: str) -> Dict[str, any]:
        """
        Find and extract projects block from CV text.
        
        Returns:
            Dict with 'found', 'start_pos', 'end_pos', 'projects_text', 'confidence'
        """
        lines = cv_text.split('\n')
        project_start = -1
        project_end = -1
        confidence = 0
        
        # Strategy 1: Look for project section headers
        for i, line in enumerate(lines):
            line_clean = line.strip().lower()
            
            # Check if this line matches project section pattern
            for pattern in self.compiled_section_patterns:
                if pattern.search(line_clean):
                    project_start = i
                    print (f"Got project at line {i} with {pattern}")
                    confidence = 0.8
                    break
            
            # Additional confidence if we find project roles/period patterns shortly after
            if project_start != -1 and i > project_start:
                if any(pattern.search(line_clean) for pattern in self.compiled_role_patterns):
                    confidence = min(confidence + 0.1, 1.0)
                if any(pattern.search(line_clean) for pattern in self.compiled_period_patterns):
                    confidence = min(confidence + 0.1, 1.0)
                
                # Look for bullet points or project descriptions
                if re.search(r'^[•\-\*]\s', line.strip()) or re.search(r'^[a-z]\.\s', line.strip()):
                    confidence = min(confidence + 0.1, 1.0)
        
        # Strategy 2: If no clear section header found, look for project-like content
        if project_start == -1:
            for i, line in enumerate(lines):
                if self._looks_like_project_content(line):
                    project_start = i
                    confidence = 0.6
                    break
        
        # Find the end of projects block
        if project_start != -1:
            project_end = self._find_projects_end(lines, project_start)
        
        # Extract projects text
        projects_text = ""
        if project_start != -1 and project_end != -1:
            projects_text = '\n'.join(lines[project_start:project_end])
        
        return {
            'found': project_start != -1,
            'start_pos': project_start,
            'end_pos': project_end,
            'projects_text': projects_text,
            'confidence': confidence,
            'line_count': project_end - project_start if project_end != -1 else 0,
            'total_lines': len(lines)
        }
    
    def _looks_like_project_content(self, line: str) -> bool:
        """Check if a line looks like project content."""
        line_clean = line.strip().lower()
        
        # Project title indicators (all caps, specific keywords)
        if (line_clean.isupper() and len(line_clean) > 5 and 
            not any(keyword in line_clean for keyword in ['education', 'language', 'domain', 'skill'])):
            return True
        
        # Project role indicators
        if any(pattern.search(line_clean) for pattern in self.compiled_role_patterns):
            return True
        
        # Project period indicators
        if any(pattern.search(line_clean) for pattern in self.compiled_period_patterns):
            return True
        
        # Responsibilities/achievements indicators
        if any(keyword in line_clean for keyword in ['responsibilities', 'achievements', 'duties', 'environment']):
            return True
        
        return False
    
    def _find_projects_end(self, lines: List[str], start_idx: int) -> int:
        """Find the end of projects block."""
        end_indicators = [
            'education', 'skills', 'certificates', 'languages', 
            'contact', 'references', 'personal', 'hobbies',
            'образование', 'навыки', 'контакты', 'личные'
        ]
        
        for i in range(start_idx + 1, len(lines)):
            line_clean = lines[i].strip().lower()
            
            # Check for next major section
            if any(indicator in line_clean for indicator in end_indicators):
                return i
            
            # Check for empty lines followed by new section
            if (line_clean == '' and i + 1 < len(lines) and 
                lines[i + 1].strip() != '' and 
                not self._looks_like_project_content(lines[i + 1])):
                return i
        
        return len(lines)


# Usage examples
if __name__ == "__main__":
    parser = ProjectBlockParser()
    try:
        parser.parse_inniwise_std("cv_lev.docx")
    except Exception as e:
        print(f"DOCX parsing error: {e}")
    
