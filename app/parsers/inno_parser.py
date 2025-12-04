"""
Main parser for InnoWise standard CV documents.
"""

import time
from typing import List
from .base_parser import BaseDocxParser, DOCX_AVAILABLE
from .models import CV, PersonalInfo, Project
from .text_normalizer import TextNormalizer
from .project_parser import InnoProjectParser

if DOCX_AVAILABLE:
    from docx import Document


class InnoStandardParser(BaseDocxParser):
    """Main parser for InnoWise standard CV documents."""
    
    def __init__(self):
        super().__init__()
        self.normalizer = TextNormalizer()
        self.project_parser = InnoProjectParser()
    
    def parse(self, file_path: str) -> CV:
        """
        Parses a standard InnoWise DOCX CV file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            CV object with parsed data
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If document structure is unexpected
        """
        self._check_file_exists(file_path)
        
        try:
            doc = Document(file_path)
            personal_info = self._parse_personal_info(doc)
            projects = self._parse_projects(doc)
            
            cv_id = f"{personal_info.name}_{int(time.time())}"
            
            return CV(
                personal_info=personal_info,
                projects=projects,
                cv_id=cv_id
            )
            
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file {file_path}: {e}")
    
    def _parse_personal_info(self, doc: Document) -> PersonalInfo:
        """Parses personal information from document."""
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        if len(paragraphs) < 3:
            raise ValueError(f"Expected at least 3 paragraphs, got {len(paragraphs)}")
        
        name = paragraphs[0]
        position_text = paragraphs[1]
        
        level, position_clean = self.normalizer.extract_position_level(position_text)
        roles = [self.normalizer.normalize(role) for role in position_clean.split('/')]
        
        # Parse table with additional information
        if not doc.tables:
            raise ValueError("Document has no tables")
        
        about_table = doc.tables[0]
        about_text = about_table.rows[0].cells[0].text
        about_lines = about_text.split('\n')
        
        education = self.normalizer.extract_between_markers(
            about_lines, "Education", "Language proficiency"
        )
        
        languages_raw = self.normalizer.extract_between_markers(
            about_lines, "Language proficiency", "Domains"
        )
        
        # Clean language entries
        languages = []
        for lang in languages_raw:
            cleaned, _ = self.normalizer.clean_language_entry(lang)
            languages.append(cleaned)
        
        domains = self.normalizer.extract_between_markers(
            about_lines, "Domains", "Certificates"
        )
        
        # Get description from second cell
        description = ""
        if len(about_table.rows[0].cells) > 1:
            description_lines = about_table.rows[0].cells[1].text.split('\n')
            description = description_lines[1] if len(description_lines) > 1 else ""
        
        return PersonalInfo(
            name=name,
            level=level,
            roles=roles,
            education=education,
            languages=languages,
            domains=domains,
            description=self.normalizer.normalize(description)
        )
    
    def _parse_projects(self, doc: Document) -> List[Project]:
        """Parses projects from the second table."""
        if len(doc.tables) < 2:
            return []
        
        projects_table = doc.tables[1]
        projects = []
        
        for row in projects_table.rows[1:]:  # Skip header
            try:
                project = self.project_parser.parse_project_row(row.cells)
                projects.append(project)
            except Exception as e:
                print(f"Warning: Failed to parse project row: {e}")
                continue
        
        return projects