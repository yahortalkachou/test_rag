import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import time

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support disabled.")

class InnoStandardPraser:
    
    def __init__(self):
        self.cv_s = {
            "metadatas": [],
            "texts": []
        }
        self.projects = []
    
    @property
    def get_metadatas(self):
        '''Get all metadatas'''
        return self.cv_s["metadatas"]

    @property
    def get_texts(self):
        '''Get all texts'''
        return self.cv_s["texts"]
     
    def _file_check(self, file_path: str):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    def _get_instance_between_keywords (self, lst: List[str], begin: str, end: str) -> List[str]:
        try:
            start = lst.index(begin) + 1
        except Exception as e:
            return []
        try:
            stop = lst.index(end)
            normalized_result = [self.normalize_string(item) for item in lst[start:stop]]
            return normalized_result
        except Exception as e:
            normalized_result = [self.normalize_string(item) for item in lst[start:]]
            return normalized_result
        
    def extract_and_clean_language_level(self, text: str) -> Tuple[str, Optional[str]]:
            pattern = r'([ABCabc])[\-\s]*(\d{1,2})\b'
    
            # Ищем совпадения
            match = re.search(pattern, text, re.IGNORECASE)
            
            if match:
                # Нашли букву уровня и цифру
                level_letter = match.group(1).upper()
                level_digit = match.group(2)
                level_full = f"{level_letter}{level_digit}"
                
                # Находим позицию начала буквы уровня
                level_start = match.start(1)
                
                # Берем часть строки до буквы уровня
                text_before = text[:level_start].strip()
                
                # Очищаем текст до уровня: убираем все не-буквенные символы в конце
                # кроме пробелов (они нам нужны для разделения слов)
                text_before_cleaned = re.sub(r'[^\w\s]+$', '', text_before)
                text_before_cleaned = re.sub(r'[\-\s]+$', '', text_before_cleaned)  # убираем дефисы/пробелы в конце
                text_before_cleaned = text_before_cleaned.strip()
                
                # Формируем результат
                result = f"{text_before_cleaned} {level_full}"
                
                return result, level_full
            
            # Если не нашли A/B/C+цифра, пробуем найти просто цифру после разделителей
            # Ищем любую цифру после не-цифровых символов в конце строки
            fallback_pattern = r'([^\d]*?)(\d{1,2})\b'
            fallback_match = re.search(fallback_pattern, text)
            
            if fallback_match:
                # Нашли цифру в конце
                text_before = fallback_match.group(1).strip()
                level_digit = fallback_match.group(2)
                
                # Очищаем текст до цифры
                text_before_cleaned = re.sub(r'[^\w\s]+$', '', text_before)
                text_before_cleaned = re.sub(r'[\-\s]+$', '', text_before_cleaned)
                text_before_cleaned = text_before_cleaned.strip()
                
                if text_before_cleaned:
                    return f"{text_before_cleaned} {level_digit}", level_digit
            
            # Если ничего не нашли, возвращаем очищенную строку
            cleaned = re.sub(r'\s+', ' ', text).strip()
            return cleaned, None
    
    def extract_and_remove_level(self,text: str) -> Tuple[Optional[str], str]:
        """
        Извлекает уровень (SENIOR, MIDDLE, JUNIOR) из строки и удаляет его.
        
        Args:
            text: Входная строка, например "SENIOR AI ENGINEER / MACHINE LEARNING ENGINEER"
            
        Returns:
            Tuple[optional уровень, очищенная строка]
        """
        level_patterns = [
            r'\b(SENIOR)\b',
            r'\b(JUNIOR)\b', 
            r'\b(MIDDLE)\b',
            r'\b(LEAD)\b',  # дополнительные возможные уровни
            r'\b(INTERN|INTERNSHIP)\b',
            r'\b(ENTRY[- ]LEVEL)\b',
            r'\b(PRINCIPAL)\b',
            r'\b(STAFF)\b',
            r'\b(SR\.)\b',  # сокращения
            r'\b(JR\.)\b'
        ]
        
        # Приводим к верхнему регистру для поиска
        upper_text = text.upper()
        found_level = None
        
        # Ищем уровень
        for pattern in level_patterns:
            match = re.search(pattern, upper_text, re.IGNORECASE)
            if match:
                found_level = match.group(1).upper()
                # Удаляем найденный уровень из оригинальной строки
                # Используем re.escape чтобы корректно обработать спецсимволы
                clean_pattern = re.compile(re.escape(match.group(1)), re.IGNORECASE)
                text = clean_pattern.sub('', text, 1)
                break
        
        # Очистка строки от лишних разделителей
        if found_level:
            # Убираем лишние пробелы и разделители
            text = re.sub(r'\s+/\s*/\s*', '/', text)  # двойные слеши
            text = re.sub(r'^\s*[/\s]+|[/\s]+\s*$', '', text)  # разделители в начале/конце
            text = re.sub(r'\s+', ' ', text).strip()  # лишние пробелы
            # Убираем одиночные слеши
            text = re.sub(r'^\s*/\s*|\s*/\s*$', '', text)
            found_level = self.normalize_string(found_level)
            
        text = self.normalize_string(text)
        return found_level, text
    
    def normalize_string(self,text: str) -> str:
        """
        Приводит строку к нижнему регистру и убирает лишние пробелы.
        
        Args:
            text: Входная строка
            
        Returns:
            Нормализованная строка
        """
        # Приводим к нижнему регистру
        normalized = text.lower()
        
        # Убираем пробелы в начале и конце
        normalized = normalized.strip()
        
        # Заменяем множественные пробелы на один (опционально)
        normalized = ' '.join(normalized.split())
        
        return normalized
    
    def _get_innowise_project_from_table (self, cells: List)->Dict:
        descriprion = cells[1].text
        project_name = cells[0].text.split('\n')
        summury = project_name[1]
        project_name = project_name[0]
        roles = self._get_instance_between_keywords(descriprion.split('\n'),"Project roles","Period")[0].split('/')
        descriprion = summury + descriprion
        
        return {
            "name": project_name,
            "description": descriprion,
            "roles" : roles,
            "CV_id" : "L1"
        }
                 
    def parse_inniwise_std (self, file_path) -> Dict[str,any]:
        """
        Extract personal and project data from .docx innowise standard file.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Extracted text as string
        """
        self._file_check(file_path)
        name = ""
        roles = []
        eduction = ""
        languages = []
        domains = []
        personal_data = ""
        projects = []
        level = ""
        
        para_texts = []
        
        try:
            doc = Document(file_path)
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_texts.append(paragraph.text)
            if len(para_texts) == 3:
                name = para_texts[0]
                #level =  para_texts[1].split(" ")[0]
                print(para_texts[1])
                level, roles = self.extract_and_remove_level(para_texts[1])
                print(roles)
                roles = roles.split("/")
                roles = [ self.normalize_string(role) for role in roles]
                #roles[0] = " ".join(roles[0].split(" ")[1:])
            #print(name,level, roles)
                            
            # Extract from std tables
            if len(doc.tables) == 3:
                #pesonal/global metadat
                about = doc.tables[0].rows[0].cells[0].text
                about = about.split('\n')
                #print (about)
                eduction = self._get_instance_between_keywords(about,"Education","Language proficiency")
                languages = self._get_instance_between_keywords(about,"Language proficiency","Domains")
                for (i,lang) in enumerate(languages):
                        cleaned = re.sub(r'[^\w\s]', '', lang)
                        
                        # Заменяем множественные пробелы на один
                        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                        languages[i] = self.normalize_string(cleaned)
                        i += 1
                
                domains = self._get_instance_between_keywords(about,"Domains","Certificates")
                cert = self._get_instance_between_keywords(about,"Certificates","axaxaxax")
                #print (f"ed: {eduction} dom: {domains} langs: {languages} certs: {cert}")
                description = doc.tables[0].rows[0].cells[1].text.split('\n')[1]
                #print(f"description\n{description}")
                #projectdata section
                
                if len(doc.tables) == 3:
                    table = doc.tables[1]
                    for row in table.rows[1:]:
                        projects.append(self._get_innowise_project_from_table(row.cells))
                
            self.cv_s["metadatas"].append({
                "CV_id": name+str(time.time_ns()),
                "name": name,
                "level": level,
                "roles": roles,
                "education": eduction,
                "languages": languages,
                "domains": domains
                #"projects": projects,
                #"description": description
            })
            self.cv_s["texts"].append(description)
            # self.global_data = {
            #     "CV_id": "L1",
            #     "name": name,
            #     "level": level,
            #     "roles": roles,
            #     "education": eduction,
            #     "languages": languages,
            #     "domains": domains,
            #     "projects": projects,
            #     "description": description
            # }
            #self.projects.append()
            return {
                "CV_id": "L1",
                "name": name,
                "level": level,
                "roles": roles,
                "education": eduction,
                "languages": languages,
                "domains": domains,
                "projects": projects,
                "description": description
            }
        
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {e}")
    
    def get_personal_data (self, n: int = 0):
        return {
                'metadata': {
                    "name": self.cv_s["metadatas"][n]["name"],
                    "level": self.cv_s["metadatas"][n]["level"],
                    "roles": self.cv_s["metadatas"][n]["roles"],
                    "education": self.cv_s["metadatas"][n]["education"],
                    "languages": self.cv_s["metadatas"][n]["languages"],
                    "domains": self.cv_s["metadatas"][n]["domains"]
                    },
                "text": self.cv_s["texts"][n]
        }
    
    # def get_projects_data (self, n: int = -1):
    #     return self.cv_s[n]["projects"]
    



if __name__ == "__main__":
    parser = InnoStandardPraser()
    try:
        parser.parse_inniwise_std("cv_lev.docx")
    except Exception as e:
        print(f"DOCX parsing error: {e}")
    
